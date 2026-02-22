"""
Utility functions for extracting text from various file formats.
"""
import os
import logging
import chardet
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'.txt', '.docx', '.pdf'}


def extract_text_from_file(uploaded_file) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text content from an uploaded file.
    
    Supports:
    - .txt files (with automatic encoding detection)
    - .docx files (Word documents)
    - .pdf files
    
    Args:
        uploaded_file: Django UploadedFile object
        
    Returns:
        Tuple[str, Optional[str]]: (extracted_text, error_message)
        - If successful: (text, None)
        - If error: (None, error_message)
    """
    if not uploaded_file:
        return None, "No file provided"
    
    # Get file extension
    file_name = uploaded_file.name
    _, ext = os.path.splitext(file_name.lower())
    
    if ext not in ALLOWED_EXTENSIONS:
        return None, f"Unsupported file format. Allowed formats: {', '.join(ALLOWED_EXTENSIONS)}"
    
    try:
        if ext == '.txt':
            return _extract_from_txt(uploaded_file)
        elif ext == '.docx':
            return _extract_from_docx(uploaded_file)
        elif ext == '.pdf':
            return _extract_from_pdf(uploaded_file)
    except Exception as e:
        logger.error(f"Error extracting text from {file_name}: {str(e)}", exc_info=True)
        return None, f"Error reading file: {str(e)}"
    
    return None, "Unknown error"


def _extract_from_txt(uploaded_file) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text from a .txt file with automatic encoding detection.
    
    Args:
        uploaded_file: Django UploadedFile object
        
    Returns:
        Tuple[str, Optional[str]]: (text, error_message)
    """
    try:
        # Read file content as bytes
        file_content = uploaded_file.read()
        
        # Detect encoding
        detected = chardet.detect(file_content)
        encoding = detected.get('encoding', 'utf-8')
        confidence = detected.get('confidence', 0)
        
        # If confidence is low, try common encodings
        if confidence < 0.7:
            for enc in ['utf-8', 'utf-8-sig', 'windows-1256', 'iso-8859-1', 'cp1252']:
                try:
                    text = file_content.decode(enc)
                    logger.info(f"Successfully decoded {uploaded_file.name} with {enc}")
                    return text, None
                except (UnicodeDecodeError, LookupError):
                    continue
        
        # Try detected encoding
        try:
            text = file_content.decode(encoding)
            logger.info(f"Successfully decoded {uploaded_file.name} with detected encoding: {encoding}")
            return text, None
        except (UnicodeDecodeError, LookupError):
            # Fallback to utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            logger.warning(f"Used UTF-8 with error replacement for {uploaded_file.name}")
            return text, None
            
    except Exception as e:
        return None, f"Error reading text file: {str(e)}"


def _extract_from_docx(uploaded_file) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text from a .docx file.
    
    Args:
        uploaded_file: Django UploadedFile object
        
    Returns:
        Tuple[str, Optional[str]]: (text, error_message)
    """
    try:
        from docx import Document
        import io
        
        # Read file content
        file_content = uploaded_file.read()
        
        # Create a file-like object from bytes
        docx_file = io.BytesIO(file_content)
        
        # Open document
        doc = Document(docx_file)
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        text = '\n'.join(text_parts)
        return text, None
        
    except ImportError:
        return None, "python-docx library is not installed. Please install it using: pip install python-docx"
    except Exception as e:
        return None, f"Error reading DOCX file: {str(e)}"


def _extract_from_pdf(uploaded_file) -> Tuple[Optional[str], Optional[str]]:
    """
    Extract text from a .pdf file.
    
    Args:
        uploaded_file: Django UploadedFile object
        
    Returns:
        Tuple[str, Optional[str]]: (text, error_message)
    """
    try:
        import pdfplumber
        import io
        
        # Read file content
        file_content = uploaded_file.read()
        
        # Create a file-like object from bytes
        pdf_file = io.BytesIO(file_content)
        
        # Extract text from all pages
        text_parts = []
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        text = '\n'.join(text_parts)
        return text, None
        
    except ImportError:
        # Try PyPDF2 as fallback
        try:
            import PyPDF2
            import io
            
            file_content = uploaded_file.read()
            pdf_file = io.BytesIO(file_content)
            
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            result = '\n'.join(text_parts)
            return result, None
            
        except ImportError:
            return None, "PDF library is not installed. Please install pdfplumber or PyPDF2: pip install pdfplumber"
    except Exception as e:
        return None, f"Error reading PDF file: {str(e)}"


def validate_file_size(uploaded_file, max_size_mb: int = 10) -> Tuple[bool, Optional[str]]:
    """
    Validate file size.
    
    Args:
        uploaded_file: Django UploadedFile object
        max_size_mb: Maximum file size in MB (default: 10MB)
        
    Returns:
        Tuple[bool, Optional[str]]: (is_valid, error_message)
    """
    if not uploaded_file:
        return False, "No file provided"
    
    max_size_bytes = max_size_mb * 1024 * 1024
    
    # Check file size
    if uploaded_file.size > max_size_bytes:
        return False, f"File size exceeds maximum allowed size of {max_size_mb}MB"
    
    return True, None


def get_file_info(uploaded_file) -> dict:
    """
    Get information about an uploaded file.
    
    Args:
        uploaded_file: Django UploadedFile object
        
    Returns:
        dict: File information including name, size, extension
    """
    if not uploaded_file:
        return {}
    
    file_name = uploaded_file.name
    _, ext = os.path.splitext(file_name.lower())
    size_mb = uploaded_file.size / (1024 * 1024)
    
    return {
        'name': file_name,
        'size': uploaded_file.size,
        'size_mb': round(size_mb, 2),
        'extension': ext,
    }

